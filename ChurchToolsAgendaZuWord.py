import logging
import tkinter
from datetime import datetime

import docx


class ChurchToolsAgendaZuWord:
    def __init__(self, api):
        self.lbx1 = None
        self.win = None
        self.api = api
        logging.basicConfig(filename='logs/main.py.log', encoding='utf-8',
                            format="%(asctime)s %(name)-10s %(levelname)-8s %(message)s",
                            level=logging.DEBUG)
        logging.debug("ChurchToolsAgendaZuWord initialized")
        self.events = []
        self.serviceGroups = self.api.get_event_masterdata(type='serviceGroups', returnAsDict=True)
        self.event_agendas = []
        self.load_events_with_agenda()
        self.create_gui()

    def load_events_with_agenda(self):
        events_temp = self.api.get_events()
        logging.debug("{} Events loaded".format(len(events_temp)))

        for event in events_temp:
            agenda = self.api.get_event_agenda(event['id'])
            if agenda is not None:
                self.event_agendas.append(agenda)
                self.events.append(event)
        logging.debug("{} Events kept because schedule exists".format(len(events_temp)))

    def create_gui(self):
        win = tkinter.Tk()
        win.title = 'Bitte Event auswählen'

        lbl1 = tkinter.Label(win, text="Die nächsten Veranstaltungen:")
        lbl1.pack()

        lbx1 = tkinter.Listbox(win, width=500)
        i = 0
        for event in self.events:
            startdate = datetime.fromisoformat(event['startDate'][:-1])
            datetext = startdate.astimezone().strftime('%a %b %d\t%H:%M')
            lbx1.insert(i, datetext + '\t' + event['name'])
            i += 1

        lbx1.pack()

        btn1 = tkinter.Button(win, text='Veranstaltung als Text umwandeln', command=self.btn1_press)
        logging.debug("GUI Elements defined")

        btn1.pack()

        logging.debug("GUI Mainloop started")

        self.win = win
        self.lbx1 = lbx1

        win.mainloop()

    def btn1_press(self):
        logging.debug("Button 1 pressed")
        if len(self.lbx1.curselection()) == 0:
            logging.info("No item selected")
            return
        else:
            event = self.events[self.lbx1.curselection()[0]]
            logging.debug("Selected event ID: {}".format(event['id']))
            # TODO #4 add FileChooser for destination path
            selected = [1]
            selectedServiceGroups = {key: value for key, value in self.serviceGroups.items() if
                                     key in selected}  # TODO 1 allow for user choice ...
            self.process_agenda(self.event_agendas[self.lbx1.curselection()[0]], serviceGroups=selectedServiceGroups)
            self.win.destroy()

    def process_agenda(self, agenda, serviceGroups, excludeBeforeEvent=True):
        logging.debug('Processing: ' + agenda['name'])

        document = docx.Document()
        heading = agenda['name']
        heading += '- Entwurf' if not agenda['isFinal'] else ''
        document.add_heading(heading)
        modifiedDate = datetime.strptime(agenda["meta"]['modifiedDate'], '%Y-%m-%dT%H:%M:%S%z')
        modifiedDate2 = modifiedDate.astimezone().strftime('%a %d.%m (%H:%M:%S)')
        document.add_paragraph("Abzug aus CT mit Änderungen bis inkl.: " + modifiedDate2)

        agenda_item = 0  # Position Argument from Event Agenda is weird therefore counting manually
        pre_event_last_item = True  # Event start is no item therefore look for change

        for item in agenda["items"]:
            if excludeBeforeEvent and item['isBeforeEvent']:
                continue

            if item['type'] == 'header':
                document.add_heading(item["title"], level=1)
                continue

            if pre_event_last_item:  # helper for event start heading which is not part of the api
                if not item['isBeforeEvent']:
                    pre_event_last_item = False
                    document.add_heading('Eventstart', level=1)

            agenda_item += 1

            title = str(agenda_item)
            title += ' ' + item["title"]

            if item['type'] == 'song':
                title += ': ' + item['song']['title']
                title += ' (' + item['song']['category'] + ')'  # TODO #5 check if fails on empty song items

            document.add_heading(title, level=2)

            responsible_list = []
            for responsible_item in item['responsible']['persons']:
                if responsible_item['person'] is not None:
                    responsible_text = responsible_item['person']['title']
                    if not responsible_item['accepted']:
                        responsible_text += ' (Angefragt)'
                else:
                    responsible_text = '?'
                responsible_text += ' ' + responsible_item['service'] + ''
                responsible_list.append(responsible_text)

            if len(item['responsible']) > 0 and len(item['responsible']['persons']) == 0:
                if len(item['responsible']['text']) > 0:
                    responsible_list.append(
                        item['responsible']['text'] + ' (Person statt Rolle in ChurchTools hinterlegt!)')

            responsible_text = ", ".join(responsible_list)
            document.add_paragraph(responsible_text)

            if item['note'] is not None and item['note'] != '':
                document.add_paragraph(item["note"])

            if len(item['serviceGroupNotes']) > 0:
                for note in item['serviceGroupNotes']:
                    if note['serviceGroupId'] in serviceGroups.keys() and len(note['note']) > 0:
                        document.add_heading("Bemerkung für {}:".format(serviceGroups[note['serviceGroupId']]['name']),
                                             level=4)
                        document.add_paragraph(note['note'])

        document.save('output/' + agenda['name'] + '.docx')
